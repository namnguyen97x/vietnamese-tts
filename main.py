import sys, os
import asyncio
import tempfile
import uuid
import requests
import webbrowser
import browser_cookie3
import json
from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                           QHBoxLayout, QTextEdit, QPushButton, QTabWidget,
                           QLabel, QComboBox, QFileDialog, QMessageBox, QDialog, QDialogButtonBox, QListWidget, QListWidgetItem, QStatusBar)
from PyQt6.QtCore import Qt, QUrl, QThread, pyqtSignal, QTimer, QByteArray, QDateTime
from PyQt6.QtNetwork import QNetworkCookie
from PyQt6.QtMultimedia import QMediaPlayer, QAudioOutput
from PyQt6.QtWebEngineWidgets import QWebEngineView
from PyQt6.QtWebEngineCore import QWebEnginePage, QWebEngineProfile
import edge_tts
from gtts import gTTS
from docx import Document
import speech_recognition as sr
from pydub import AudioSegment
from pydub.utils import which
import subprocess
import sounddevice as sd
import numpy as np
import wave
import docx
import fitz
from PyQt6.QtGui import QGuiApplication

def get_bin_dir():
    if getattr(sys, 'frozen', False):
        exe_dir = os.path.dirname(sys.executable)
        return os.path.join(exe_dir, 'bin')
    else:
        return os.path.join(os.path.dirname(os.path.abspath(__file__)), 'bin')

bin_dir = get_bin_dir()
os.environ["PATH"] += os.pathsep + bin_dir
ffmpeg_path = os.path.join(bin_dir, 'ffmpeg.exe')
ffprobe_path = os.path.join(bin_dir, 'ffprobe.exe')
print("[DEBUG] ffmpeg_path:", ffmpeg_path)
print("[DEBUG] ffprobe_path:", ffprobe_path)
if os.path.exists(ffmpeg_path):
    print(f"[INFO] Đã tìm thấy ffmpeg tại: {ffmpeg_path}")
    AudioSegment.converter = ffmpeg_path
    AudioSegment.ffmpeg = ffmpeg_path
else:
    print(f"[WARNING] Không tìm thấy ffmpeg tại: {ffmpeg_path}. Hãy kiểm tra lại tên file và vị trí!")
if os.path.exists(ffprobe_path):
    print(f"[INFO] Đã tìm thấy ffprobe tại: {ffprobe_path}")
    AudioSegment.ffprobe = ffprobe_path
else:
    print(f"[WARNING] Không tìm thấy ffprobe tại: {ffprobe_path}. Hãy kiểm tra lại tên file và vị trí!")

class GeminiWebPage(QWebEnginePage):
    popup_url_found = pyqtSignal(QUrl)

    def createWindow(self, _type):
        popup_page = GeminiWebPage(self.profile(), self)
        popup_page.urlChanged.connect(self.popup_url_found)
        return popup_page

class EdgeTTSWorker(QThread):
    voices_loaded = pyqtSignal(list)
    finished = pyqtSignal()
    error = pyqtSignal(str)

    def __init__(self, text=None, voice=None, output_file=None, load_voices_only=False):
        super().__init__()
        self.text = text
        self.voice = voice
        self.output_file = output_file
        self.load_voices_only = load_voices_only

    def run(self):
        try:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            
            if self.load_voices_only:
                voices = loop.run_until_complete(edge_tts.list_voices())
                self.voices_loaded.emit(voices)
            else:
                if not self.text or not self.voice:
                    raise ValueError("Text và Voice không được để trống khi tạo âm thanh.")
                communicate = edge_tts.Communicate(self.text, self.voice)
                loop.run_until_complete(communicate.save(self.output_file))
                self.finished.emit()
        except Exception as e:
            self.error.emit(str(e))

class GttsWorker(QThread):
    finished = pyqtSignal(str)
    error = pyqtSignal(str)

    def __init__(self, text, lang):
        super().__init__()
        self.text = text
        self.lang = lang

    def run(self):
        try:
            tts = gTTS(self.text, lang=self.lang)
            temp_dir = tempfile.gettempdir()
            timestamp = QDateTime.currentDateTime().toString("yyyyMMdd_hhmmss")
            output_path = os.path.join(temp_dir, f"gtts_{timestamp}_{uuid.uuid4().hex[:6]}.mp3")
            tts.save(output_path)
            self.finished.emit(output_path)
        except Exception as e:
            self.error.emit(str(e))

class STTFileWorker(QThread):
    result = pyqtSignal(str)
    error = pyqtSignal(str)
    def __init__(self, file_path, ffmpeg_path):
        super().__init__()
        self.file_path = file_path
        self.ffmpeg_path = ffmpeg_path
    def run(self):
        import speech_recognition as sr
        import tempfile, os, sys, subprocess
        temp_wav = None
        try:
            creationflags = subprocess.CREATE_NO_WINDOW if sys.platform == "win32" else 0
            tmp_wav = tempfile.NamedTemporaryFile(suffix='.wav', delete=False)
            tmp_wav.close()
            cmd = [self.ffmpeg_path, '-y', '-i', self.file_path, '-acodec', 'pcm_s16le', '-ac', '1', '-ar', '16000', tmp_wav.name]
            subprocess.run(cmd, check=True, creationflags=creationflags)
            temp_wav = tmp_wav.name
            recognizer = sr.Recognizer()
            with sr.AudioFile(temp_wav) as source:
                audio = recognizer.record(source)
            text = recognizer.recognize_google(audio, language='vi-VN')
            self.result.emit(text)
        except Exception as e:
            self.error.emit(f"Lỗi nhận diện: {e}")
        finally:
            if temp_wav and os.path.exists(temp_wav):
                os.remove(temp_wav)

class TTSApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Vietnamese TTS")
        self.resize(1000, 700)
        self.temp_dir = tempfile.mkdtemp(prefix="tts_app_")
        self.temp_files = []
        
        self.audio_output = QAudioOutput()
        self.player = QMediaPlayer()
        self.player.setAudioOutput(self.audio_output)
        self.player.errorOccurred.connect(self.player_error)

        self.statusBar = QStatusBar()
        self.setStatusBar(self.statusBar)
        self.init_ui()

    def init_ui(self):
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QHBoxLayout(central_widget)

        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)
        self.tabs = QTabWidget()
        self.tabs.currentChanged.connect(self.on_tab_changed)
        
        edge_tab = QWidget()
        self.init_edge_tts_tab(edge_tab)
        self.tabs.addTab(edge_tab, "Edge TTS")

        google_tab = QWidget()
        self.init_gtts_tab(google_tab)
        self.tabs.addTab(google_tab, "Google TTS (gtts)")

        stt_tab = QWidget()
        self.init_stt_tab(stt_tab)
        self.tabs.addTab(stt_tab, "Speech to Text (STT)")

        ai_chat_tab = QWidget()
        ai_chat_layout = QVBoxLayout(ai_chat_tab)
        self.ai_subtabs = QTabWidget()

        # Tab Gemini
        self.gemini_inner_tab = QWidget()
        self.init_gemini_tab(self.gemini_inner_tab)
        self.ai_subtabs.addTab(self.gemini_inner_tab, "Gemini")

        # Tab ChatGPT
        self.chatgpt_inner_tab = QWidget()
        self.init_chatgpt_tab(self.chatgpt_inner_tab)
        self.ai_subtabs.addTab(self.chatgpt_inner_tab, "ChatGPT")

        ai_chat_layout.addWidget(self.ai_subtabs)
        self.tabs.addTab(ai_chat_tab, "AI chat")

        left_layout.addWidget(self.tabs)
        main_layout.addWidget(left_widget, 2)

        right_widget = QWidget()
        right_layout = QVBoxLayout(right_widget)
        right_layout.addWidget(QLabel("Danh sách file tạm:"))
        self.temp_file_list_widget = QListWidget()
        self.temp_file_list_widget.itemClicked.connect(self.on_temp_file_selected)
        right_layout.addWidget(self.temp_file_list_widget)
        
        audio_btn_layout = QHBoxLayout()
        self.play_button = QPushButton("Phát")
        self.play_button.setEnabled(False)
        self.play_button.clicked.connect(self.play_selected_file)
        self.save_file_button = QPushButton("Lưu file")
        self.save_file_button.setEnabled(False)
        self.save_file_button.clicked.connect(self.save_selected_file)
        self.delete_file_button = QPushButton("Xóa file tạm")
        self.delete_file_button.clicked.connect(self.delete_selected_file)
        audio_btn_layout.addWidget(self.play_button)
        audio_btn_layout.addWidget(self.save_file_button)
        audio_btn_layout.addWidget(self.delete_file_button)
        right_layout.addLayout(audio_btn_layout)
        main_layout.addWidget(right_widget, 1)

    def closeEvent(self, event):
        self.gemini_view.setPage(None)
        self.gemini_page.deleteLater()
        self.gemini_profile.deleteLater()
        super().closeEvent(event)

    def on_tab_changed(self, index):
        if self.player.playbackState() == QMediaPlayer.PlaybackState.PlayingState:
            self.player.stop()
            self.play_button.setText("Phát")

    def add_temp_file(self, name, path, file_type):
        self.temp_files.append({'name': name, 'path': path, 'type': file_type})
        self.temp_file_list_widget.addItem(QListWidgetItem(name))

    def on_temp_file_selected(self):
        idx = self.temp_file_list_widget.currentRow()
        if idx < 0: return
        
        file_info = self.temp_files[idx]
        if file_info['type'] == 'audio':
            self.play_button.setEnabled(True)
            self.save_file_button.setText("Lưu file âm thanh")
            self.save_file_button.setEnabled(True)
            self.play_audio(file_info['path'])
        elif file_info['type'] == 'image':
            self.play_button.setEnabled(False)
            self.save_file_button.setText("Lưu file ảnh")
            self.save_file_button.setEnabled(True)
            # Stop audio if it was playing and an image is selected
            if self.player.playbackState() == QMediaPlayer.PlaybackState.PlayingState:
                self.player.stop()
                self.play_button.setText("Phát")

    def play_selected_file(self):
        idx = self.temp_file_list_widget.currentRow()
        if idx < 0: return
        
        file_info = self.temp_files[idx]
        if file_info['type'] != 'audio': return # Do not play non-audio files
        
        path = file_info['path']
        if self.player.source() == QUrl.fromLocalFile(path) and self.player.playbackState() == QMediaPlayer.PlaybackState.PlayingState:
            self.player.pause()
            self.play_button.setText("Phát")
        elif self.player.source() == QUrl.fromLocalFile(path) and self.player.playbackState() == QMediaPlayer.PlaybackState.PausedState:
            self.player.play()
            self.play_button.setText("Tạm dừng")
        else:
            self.play_audio(path)
            
    def play_audio(self, path):
        if path:
            self.player.setSource(QUrl.fromLocalFile(path))
            self.player.play()
            self.play_button.setText("Tạm dừng")

    def save_selected_file(self):
        idx = self.temp_file_list_widget.currentRow()
        if idx < 0: return
        
        file_info = self.temp_files[idx]
        path = file_info['path']
        name = file_info['name']
        ext = os.path.splitext(path)[1]
        
        if file_info['type'] == 'audio':
            filter = f"Audio Files (*{ext})"
        elif file_info['type'] == 'image':
            filter = "Image Files (*.png *.jpg *.jpeg *.bmp *.gif)"
        else:
            filter = "All Files (*)"

        save_name_suggestion = os.path.splitext(name)[0] + ext
        file_path, _ = QFileDialog.getSaveFileName(self, "Lưu file", save_name_suggestion, filter)
        
        if file_path:
            try:
                import shutil
                shutil.copy(path, file_path)
                self.statusBar.showMessage(f"Đã lưu file thành công tới: {file_path}", 5000)
            except Exception as e:
                QMessageBox.warning(self, "Lỗi", f"Không thể lưu file: {e}")

    def delete_selected_file(self):
        idx = self.temp_file_list_widget.currentRow()
        if idx < 0: return
        
        file_info = self.temp_files.pop(idx)
        self.temp_file_list_widget.takeItem(idx)
        try:
            os.remove(file_info['path'])
        except OSError:
            pass
        
        if not self.temp_files:
            self.play_button.setText("Phát")
            self.play_button.setEnabled(False)
            self.save_file_button.setText("Lưu file")
            self.save_file_button.setEnabled(False)

    def player_error(self, error, error_string):
        QMessageBox.critical(self, "Lỗi Media Player", f"Gặp lỗi khi phát file:\n{error_string}")
        self.play_button.setText("Phát")
        
    def on_tts_error(self, error_message):
        self.statusBar.showMessage(f"Lỗi khi tạo âm thanh: {error_message}", 5000)
        self.read_aloud_button.setEnabled(True)
        self.gtts_speak_button.setEnabled(True)

    # --- Edge TTS Tab ---
    def init_edge_tts_tab(self, tab):
        layout = QVBoxLayout(tab)
        self.text_input = QTextEdit()
        self.text_input.setPlaceholderText("Nhập văn bản cần chuyển đổi tại đây...")
        layout.addWidget(self.text_input)

        voice_layout = QHBoxLayout()
        voice_layout.addWidget(QLabel("Chọn giọng đọc:"))
        self.voice_combo = QComboBox()
        voice_layout.addWidget(self.voice_combo)
        layout.addLayout(voice_layout)

        import_button = QPushButton("Nhập File Văn Bản (.txt, .pdf, .docx)")
        import_button.clicked.connect(self.import_text_file)
        layout.addWidget(import_button)

        self.loading_label = QLabel("Đang tải danh sách giọng đọc...")
        self.loading_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(self.loading_label)

        button_layout = QHBoxLayout()
        self.read_aloud_button = QPushButton("Đọc thử")
        self.read_aloud_button.setEnabled(False)
        self.read_aloud_button.clicked.connect(self.generate_speech)
        button_layout.addWidget(self.read_aloud_button)
        layout.addLayout(button_layout)
        
        self.voice_loader_worker = EdgeTTSWorker(load_voices_only=True)
        self.voice_loader_worker.voices_loaded.connect(self.populate_voices)
        self.voice_loader_worker.error.connect(self.on_voice_load_error)
        self.voice_loader_worker.start()
        
    def generate_speech(self):
        text = self.text_input.toPlainText().strip()
        if not text:
            return
        voice = self.voice_combo.currentData()
        if not voice:
            QMessageBox.warning(self, "Lỗi", "Giọng đọc chưa được chọn hoặc không hợp lệ.")
            return
            
        filename = f"edge_tts_{uuid.uuid4().hex[:8]}.mp3"
        output_file = os.path.join(self.temp_dir, filename)
        
        self.statusBar.showMessage("Đang xử lý, vui lòng chờ...")
        self.read_aloud_button.setEnabled(False)

        self.edge_worker = EdgeTTSWorker(text=text, voice=voice, output_file=output_file)
        self.edge_worker.finished.connect(lambda: self.on_edge_tts_finished(output_file))
        self.edge_worker.error.connect(self.on_tts_error)
        self.edge_worker.start()

    def on_edge_tts_finished(self, path):
        self.statusBar.showMessage(f"Đã tạo file thành công: {os.path.basename(path)}", 5000)
        self.read_aloud_button.setEnabled(True)
        self.add_temp_file(os.path.basename(path), path, 'audio')
        self.play_audio(path)

    def populate_voices(self, voices):
        self.loading_label.setText("")
        filtered_voices = [v for v in voices if v.get('Locale','').startswith('vi-') or 'Multilingual' in v.get('ShortName','')]
        if not filtered_voices:
            self.loading_label.setText("Không tìm thấy giọng đọc phù hợp.")
            return
            
        def sort_key(v):
            return (0, v['ShortName']) if 'hoaimy' in v['ShortName'].lower() or 'namminh' in v['ShortName'].lower() else (1, v['ShortName'])
        filtered_voices.sort(key=sort_key)
        
        for voice in filtered_voices:
            self.voice_combo.addItem(f"{voice.get('ShortName')} - {voice.get('FriendlyName')}", voice.get('ShortName'))
        self.read_aloud_button.setEnabled(True)

    def on_voice_load_error(self, msg):
        self.loading_label.setText(f"Lỗi tải giọng đọc: {msg}")
        self.loading_label.setStyleSheet("color: red;")

    def import_text_file(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Chọn file văn bản", "", "All Files (*);;Text Files (*.txt);;PDF Files (*.pdf);;Word Documents (*.docx)")
        if not file_path: return
        content = ""
        try:
            if file_path.lower().endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as f: content = f.read()
            elif file_path.lower().endswith('.pdf'):
                with fitz.open(file_path) as doc: content = "".join(page.get_text() for page in doc)
            elif file_path.lower().endswith('.docx'):
                doc = docx.Document(file_path)
                content = "\n".join(para.text for para in doc.paragraphs)
            self.text_input.setPlainText(content)
            self.statusBar.showMessage(f"Đã tải thành công: {os.path.basename(file_path)}", 5000)
        except Exception as e:
            QMessageBox.critical(self, "Lỗi Đọc File", f"Không thể đọc file.\nLỗi: {e}")

    # --- Google TTS Tab ---
    def init_gtts_tab(self, tab):
        layout = QVBoxLayout(tab)
        self.gtts_text_input = QTextEdit()
        layout.addWidget(QLabel("Nhập văn bản:"))
        layout.addWidget(self.gtts_text_input)
        
        lang_layout = QHBoxLayout()
        lang_layout.addWidget(QLabel("Ngôn ngữ:"))
        self.gtts_lang_combo = QComboBox()
        self.gtts_lang_combo.addItems(["vi", "en", "fr", "de", "es"])
        lang_layout.addWidget(self.gtts_lang_combo)
        layout.addLayout(lang_layout)

        self.gtts_speak_button = QPushButton("Tạo và Đọc thử")
        self.gtts_speak_button.clicked.connect(self.speak_gtts)
        layout.addWidget(self.gtts_speak_button)
        
    def speak_gtts(self):
        text = self.gtts_text_input.toPlainText().strip()
        if not text: return
        lang = self.gtts_lang_combo.currentText()
        
        self.gtts_speak_button.setEnabled(False)
        self.statusBar.showMessage("Đang xử lý với Google TTS...")
        self.gtts_worker = GttsWorker(text, lang)
        self.gtts_worker.finished.connect(self.on_gtts_finished)
        self.gtts_worker.error.connect(self.on_tts_error)
        self.gtts_worker.start()

    def on_gtts_finished(self, path):
        name = os.path.basename(path)
        self.statusBar.showMessage(f"Google TTS tạo file thành công: {name}", 5000)
        self.gtts_speak_button.setEnabled(True)
        self.add_temp_file(name, path, 'audio')
        self.play_audio(path)

    # --- STT Tab ---
    def init_stt_tab(self, tab):
        layout = QVBoxLayout(tab)
        self.stt_result = QTextEdit()
        self.stt_result.setPlaceholderText("Kết quả nhận diện sẽ hiển thị ở đây...")
        layout.addWidget(self.stt_result)

        btn_layout = QHBoxLayout()
        self.stt_file_btn = QPushButton("Nhận diện từ file")
        self.stt_file_btn.clicked.connect(self.stt_from_file)
        btn_layout.addWidget(self.stt_file_btn)
        
        self.stt_save_btn = QPushButton("Lưu văn bản...")
        self.stt_save_btn.clicked.connect(self.save_stt_text)
        btn_layout.addWidget(self.stt_save_btn)
        layout.addLayout(btn_layout)

    def stt_from_file(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Chọn file âm thanh", "", "Audio Files (*.wav *.mp3 *.m4a *.ogg)")
        if not file_path: return
        self.stt_result.setPlainText("Đang xử lý file âm thanh...")
        self.stt_file_btn.setEnabled(False)
        self.stt_worker = STTFileWorker(file_path, ffmpeg_path)
        self.stt_worker.result.connect(self.on_stt_result)
        self.stt_worker.error.connect(self.on_stt_error)
        self.stt_worker.finished.connect(lambda: self.stt_file_btn.setEnabled(True))
        self.stt_worker.start()

    def on_stt_result(self, text):
        self.stt_result.setPlainText(text)

    def on_stt_error(self, msg):
        self.stt_result.setPlainText(msg)

    def save_stt_text(self):
        text = self.stt_result.toPlainText().strip()
        if not text: return
        file_path, _ = QFileDialog.getSaveFileName(self, "Lưu văn bản", "stt_result.txt", "Text Files (*.txt);;Word Documents (*.docx)")
        if not file_path: return
        try:
            if file_path.endswith('.docx'):
                doc = docx.Document()
                doc.add_paragraph(text)
                doc.save(file_path)
            else:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(text)
            self.statusBar.showMessage("Lưu file thành công!", 3000)
        except Exception as e:
            QMessageBox.critical(self, "Lỗi Lưu File", f"Không thể lưu file.\nLỗi: {e}")

    # --- Gemini Tab ---
    def init_gemini_tab(self, tab):
        layout = QVBoxLayout(tab)
        
        self.gemini_view = QWebEngineView()
        self.gemini_profile = QWebEngineProfile("gemini_profile", self.gemini_view)
        self.gemini_profile.setPersistentCookiesPolicy(QWebEngineProfile.PersistentCookiesPolicy.AllowPersistentCookies)
        self.gemini_page = GeminiWebPage(self.gemini_profile, self.gemini_view)
        self.gemini_view.setPage(self.gemini_page)
        
        control_layout = QHBoxLayout()
        self.gemini_import_cookie_btn = QPushButton("🍪 Nhập Cookie")
        self.gemini_import_cookie_btn.clicked.connect(self.import_gemini_cookies)
        control_layout.addWidget(self.gemini_import_cookie_btn)
        self.gemini_refresh_btn = QPushButton("Làm mới")
        self.gemini_refresh_btn.clicked.connect(self.gemini_view.reload)
        control_layout.addWidget(self.gemini_refresh_btn)
        
        layout.addLayout(control_layout)
        layout.addWidget(self.gemini_view)
        
        copy_tts_btn = QPushButton("Sao chép vào TTS")
        copy_tts_btn.clicked.connect(self.extract_and_convert_gemini_text)
        layout.addWidget(copy_tts_btn)
        
        self.gemini_profile.downloadRequested.connect(self.handle_download)

        self.gemini_view.load(QUrl("https://gemini.google.com"))

    def import_gemini_cookies(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Chọn file Cookie Export", "", "JSON Files (*.json)")
        if not file_path: return
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                cookies_data = json.load(f)
            
            cookie_store = self.gemini_page.profile().cookieStore()
            # Xóa tất cả các cookie cũ để đảm bảo không bị xung đột
            cookie_store.deleteAllCookies()

            count = 0
            for cookie in cookies_data:
                if 'google.com' not in cookie.get('domain', ''): continue
                
                # Tạo QNetworkCookie từ dữ liệu JSON
                q_cookie = QNetworkCookie(
                    cookie.get('name', '').encode(),
                    cookie.get('value', '').encode()
                )
                q_cookie.setDomain(cookie.get('domain', ''))
                q_cookie.setPath(cookie.get('path', '/'))

                if 'expirationDate' in cookie and cookie['expirationDate']:
                    q_cookie.setExpirationDate(QDateTime.fromSecsSinceEpoch(int(cookie['expirationDate'])))
                
                q_cookie.setHttpOnly(cookie.get('httpOnly', False))
                q_cookie.setSecure(cookie.get('secure', False))

                # Thuộc tính sameSite không được hỗ trợ trong phiên bản PyQt6 này nên đã bị loại bỏ

                cookie_store.setCookie(q_cookie)
                count += 1
            
            self.statusBar.showMessage(f"Đã nhập {count} cookie. Đang tải lại...", 3000)
            # Dùng QTimer để chờ một chút cho cookie store xử lý xong trước khi reload
            QTimer.singleShot(1000, self.gemini_view.reload)

        except Exception as e:
            QMessageBox.critical(self, "Lỗi Nhập Cookie", f"Không thể xử lý file cookie.\nLỗi: {e}")

    def extract_and_convert_gemini_text(self):
        self.statusBar.showMessage("Đang trích xuất văn bản từ Gemini...", 3000)
        js_code = "Array.from(document.querySelectorAll('.markdown')).pop().innerText;"
        self.gemini_view.page().runJavaScript(js_code, self.on_gemini_text_extracted)

    def on_gemini_text_extracted(self, text):
        if text and text.strip():
            self.text_input.setPlainText(text.strip())
            self.tabs.setCurrentIndex(0) # Chuyển qua tab Edge TTS
            self.statusBar.showMessage("Đã sao chép văn bản vào tab Edge TTS!", 5000)
        else:
            self.statusBar.showMessage("Không tìm thấy nội dung trả lời của AI.", 5000)

    def handle_download(self, download_item):
        try:
            save_dir = os.path.join(self.temp_dir, "images")
            os.makedirs(save_dir, exist_ok=True)
            filename = download_item.url().fileName()
            if not filename:
                filename = f"download_{uuid.uuid4().hex[:6]}.img"
            
            save_path = os.path.join(save_dir, filename)
            
            download_item.setDownloadDirectory(save_dir)
            download_item.setDownloadFileName(filename)
            download_item.accept()
            
            # Thêm file vào danh sách ngay lập tức (workaround cho lỗi 'finished' signal)
            self.add_temp_file(filename, save_path, 'image')
            self.statusBar.showMessage(f"Đang tải về: {filename}", 3000)
        except Exception as e:
            self.statusBar.showMessage(f"Lỗi khi bắt đầu tải: {e}", 5000)

    def init_chatgpt_tab(self, tab):
        layout = QVBoxLayout(tab)
        self.chatgpt_view = QWebEngineView()
        self.chatgpt_profile = QWebEngineProfile("chatgpt_profile", self.chatgpt_view)
        self.chatgpt_profile.setPersistentCookiesPolicy(QWebEngineProfile.PersistentCookiesPolicy.AllowPersistentCookies)
        self.chatgpt_page = GeminiWebPage(self.chatgpt_profile, self.chatgpt_view)
        self.chatgpt_view.setPage(self.chatgpt_page)

        control_layout = QHBoxLayout()
        self.chatgpt_import_cookie_btn = QPushButton("🍪 Nhập Cookie")
        self.chatgpt_import_cookie_btn.clicked.connect(self.import_chatgpt_cookies)
        control_layout.addWidget(self.chatgpt_import_cookie_btn)
        self.chatgpt_refresh_btn = QPushButton("Làm mới")
        self.chatgpt_refresh_btn.clicked.connect(self.chatgpt_view.reload)
        control_layout.addWidget(self.chatgpt_refresh_btn)
        
        layout.addLayout(control_layout)
        layout.addWidget(self.chatgpt_view)
        
        copy_tts_btn = QPushButton("Sao chép vào TTS")
        copy_tts_btn.clicked.connect(self.extract_and_convert_chatgpt_text)
        layout.addWidget(copy_tts_btn)
        self.chatgpt_profile.downloadRequested.connect(self.handle_download)

        self.chatgpt_view.load(QUrl("https://chatgpt.com/"))

    def import_chatgpt_cookies(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Chọn file Cookie Export", "", "JSON Files (*.json)")
        if not file_path: return
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                cookies_data = json.load(f)
            cookie_store = self.chatgpt_page.profile().cookieStore()
            cookie_store.deleteAllCookies()
            count = 0
            for cookie in cookies_data:
                if 'openai.com' not in cookie.get('domain', '') and 'chatgpt.com' not in cookie.get('domain', ''): continue
                q_cookie = QNetworkCookie(
                    cookie.get('name', '').encode(),
                    cookie.get('value', '').encode()
                )
                q_cookie.setDomain(cookie.get('domain', ''))
                q_cookie.setPath(cookie.get('path', '/'))
                if 'expirationDate' in cookie and cookie['expirationDate']:
                    q_cookie.setExpirationDate(QDateTime.fromSecsSinceEpoch(int(cookie['expirationDate'])))
                q_cookie.setHttpOnly(cookie.get('httpOnly', False))
                q_cookie.setSecure(cookie.get('secure', False))
                cookie_store.setCookie(q_cookie)
                count += 1
            self.statusBar.showMessage(f"Đã nhập {count} cookie. Đang tải lại...", 3000)
            QTimer.singleShot(1000, self.chatgpt_view.reload)
        except Exception as e:
            QMessageBox.critical(self, "Lỗi Nhập Cookie", f"Không thể xử lý file cookie.\nLỗi: {e}")

    def extract_and_convert_chatgpt_text(self):
        self.statusBar.showMessage("Đang trích xuất văn bản từ ChatGPT...", 3000)
        js_code = "Array.from(document.querySelectorAll('.markdown')).pop().innerText;"
        self.chatgpt_view.page().runJavaScript(js_code, self.on_chatgpt_text_extracted)

    def on_chatgpt_text_extracted(self, text):
        if text and text.strip():
            self.text_input.setPlainText(text.strip())
            self.tabs.setCurrentIndex(0) # Chuyển qua tab Edge TTS
            self.statusBar.showMessage("Đã sao chép văn bản vào tab Edge TTS!", 5000)
        else:
            self.statusBar.showMessage("Không tìm thấy nội dung trả lời của AI.", 5000)

if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = TTSApp()
    window.show()
    sys.exit(app.exec()) 